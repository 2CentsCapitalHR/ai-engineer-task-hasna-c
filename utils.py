import os
import glob
from typing import List, Dict, Any
from sentence_transformers import SentenceTransformer
import numpy as np
import openai
from docx import Document
import json
import uuid
from dotenv import load_dotenv

# Load environment variables from .env
load_dotenv()

# Default checklist
DEFAULT_CHECKLIST = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors"
    ]
}

def load_reference_corpus(ref_folder: str) -> List[Dict[str, Any]]:
    corpus = []
    if not os.path.isdir(ref_folder):
        return corpus
    txt_files = glob.glob(os.path.join(ref_folder, "*.txt"))
    idc = 0
    for p in txt_files:
        with open(p, "r", encoding="utf-8") as f:
            text = f.read()
        paras = [para.strip() for para in text.split("\n\n") if para.strip()]
        for para in paras:
            corpus.append({"id": f"{os.path.basename(p)}::{idc}", "text": para, "meta": {"source": p}})
            idc += 1
    return corpus

def build_vector_index(corpus):
    model = SentenceTransformer("all-MiniLM-L6-v2")
    texts = [c["text"] for c in corpus]
    embeddings = model.encode(texts, convert_to_numpy=True)
    return {"embeddings": embeddings, "ids": [c["id"] for c in corpus], "model": model}

def retrieve_relevant_passages(doc_path: str, index_obj, corpus, top_k=3):
    text = extract_docx_text(doc_path)
    model = index_obj["model"]
    q_emb = model.encode([text], convert_to_numpy=True)[0]
    sims = np.dot(index_obj["embeddings"], q_emb)
    top_idx = np.argsort(sims)[-top_k:][::-1]
    return [corpus[i] for i in top_idx]

def extract_docx_text(path: str) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def analyze_documents_with_llm(doc_path: str, retrieved_passages: List[Dict], checklist: Dict = DEFAULT_CHECKLIST) -> Dict:
    openai.api_key = os.getenv("OPENAI_API_KEY")
    text = extract_docx_text(doc_path)
    prompt = f"""
You are an ADGM compliance assistant. Analyze the document below and the retrieved reference text.

Document:
<<<
{text[:3000]}
>>>

References:
"""
    for p in retrieved_passages:
        prompt += f"- {p['text'][:500]}\n"

    prompt += """
Output a JSON with:
process, documents_uploaded, required_documents, missing_document, issues_found.
issues_found = list of objects: document, section, issue, severity, suggestion.
"""

    resp = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that outputs JSON only."},
            {"role": "user", "content": prompt}
        ],
        temperature=0,
        max_tokens=800
    )

    try:
        return json.loads(resp.choices[0].message.content.strip())
    except:
        return {
            "process": "Unknown",
            "documents_uploaded": 1,
            "required_documents": len(checklist.get("Company Incorporation", [])),
            "missing_document": None,
            "issues_found": []
        }

def annotate_docx_with_findings(doc_path: str, issues: List[Dict], out_suffix="_reviewed") -> str:
    doc = Document(doc_path)
    for iss in issues:
        locator = iss.get("section") or ""
        inserted = False
        for p in doc.paragraphs:
            if locator and locator in p.text:
                p.add_run(f" [COMMENT: {iss.get('issue')} | Suggestion: {iss.get('suggestion')}]")
                inserted = True
                break
        if not inserted:
            doc.add_paragraph(f"[COMMENT on {iss.get('document')}: {iss.get('issue')} | Suggestion: {iss.get('suggestion')}]")
    out_path = doc_path.replace(".docx", f"{out_suffix}.docx")
    if os.path.exists(out_path):
        out_path = doc_path.replace(".docx", f"{out_suffix}_{uuid.uuid4().hex[:6]}.docx")
    doc.save(out_path)
    return out_path
